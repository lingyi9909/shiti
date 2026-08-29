from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from PIL import Image

from question_builder.parser.docx.body import parse_docx


def _add_numbering(paragraph: object, num_id: int, ilvl: int) -> None:
    p = paragraph._p  # type: ignore[attr-defined]
    p_pr = p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), str(ilvl))
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(level)
    num_pr.append(num)
    p_pr.append(num_pr)


def _build_fixture(path: Path) -> None:
    image_path = path.with_suffix(".png")
    Image.new("RGB", (8, 8), "white").save(image_path)

    document = Document()
    first = document.add_paragraph("First question")
    _add_numbering(first, 5, 0)
    mixed = document.add_paragraph()
    mixed.add_run("before ")
    mixed.add_run().add_picture(str(image_path), width=Inches(0.1))
    mixed.add_run(" after")
    document.add_table(rows=1, cols=1).cell(0, 0).text = "table boundary"
    document.add_paragraph("Last paragraph").add_run().add_break(WD_BREAK.LINE)
    document.save(path)


def test_basic_document_ir_preserves_exact_body_and_run_child_order(tmp_path: Path) -> None:
    source = tmp_path / "basic_ordered.docx"
    assets = tmp_path / "assets"
    _build_fixture(source)

    parsed = parse_docx(source, assets)

    assert parsed.source_file == source.name
    assert parsed.document_id == f"doc_{parsed.source_sha256[:16]}"
    assert [block.type for block in parsed.blocks] == [
        "paragraph",
        "paragraph",
        "image",
        "paragraph",
        "table",
        "paragraph",
    ]
    assert [block.raw_text for block in parsed.blocks] == [
        "First question",
        "before ",
        "",
        " after",
        "table boundary",
        "Last paragraph\n",
    ]
    assert parsed.blocks[0].numbering is not None
    assert parsed.blocks[0].numbering["resolved_label"] == "1."
    image = parsed.blocks[2]
    assert image.relationship_id
    assert image.metadata["asset_id"].startswith("img_")
    assert image.metadata["asset_sha256"]
    materialized = assets / image.metadata["asset_filename"]
    assert materialized.is_file()
    assert materialized.read_bytes() == source.with_suffix(".png").read_bytes()
    assert [block.order for block in parsed.blocks] == list(range(len(parsed.blocks)))
    assert all(block.source_xml_path for block in parsed.blocks)
